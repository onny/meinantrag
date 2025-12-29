#!/usr/bin/env python3
"""
MeinAntrag - A web application to generate prefilled government requests
"""

import falcon
import json
import requests
from urllib.parse import urlencode, parse_qs
import os
import sys
from jinja2 import Environment, FileSystemLoader
import google.generativeai as genai
import re
from io import BytesIO
from datetime import datetime
try:
	from docx import Document
	from docx.shared import Pt, Inches
	from docx.enum.text import WD_ALIGN_PARAGRAPH
	DOCX_AVAILABLE = True
except ImportError:
	DOCX_AVAILABLE = False

SITE_BASE_URL = os.environ.get('MEINANTRAG_BASE_URL', 'http://localhost:8000')

class BaseTemplateResource:
	"""Base class for resources that need template rendering"""
	
	def _get_template_dir(self):
		"""Get the template directory path, handling both development and installed environments"""
		# Allow overriding via environment variable (for packaged deployments)
		env_dir = os.environ.get('MEINANTRAG_TEMPLATES_DIR')
		if env_dir and os.path.exists(env_dir):
			return env_dir

		# Get the directory where this script is located
		script_dir = os.path.dirname(os.path.abspath(__file__))
		
		# Try development templates first
		dev_template_dir = os.path.join(script_dir, 'templates')
		if os.path.exists(dev_template_dir):
			return dev_template_dir
		
		# Try to find templates relative to the executable
		try:
			# If we're running from a Nix store, look for templates in share/meinantrag
			if '/nix/store/' in script_dir:
				# Go up from bin to share/meinantrag/templates
				share_dir = os.path.join(script_dir, '..', 'share', 'meinantrag', 'templates')
				if os.path.exists(share_dir):
					return share_dir
				
				# Alternative: look for templates in the same store path
				store_root = script_dir.split('/nix/store/')[1].split('/')[0]
				store_path = f"/nix/store/{store_root}"
				alt_share_dir = os.path.join(store_path, 'share', 'meinantrag', 'templates')
				if os.path.exists(alt_share_dir):
					return alt_share_dir
		except Exception:
			pass
		
		# Last resort: try to find any templates directory
		for root, dirs, files in os.walk('/nix/store'):
			if 'templates' in dirs and 'index.html' in os.listdir(os.path.join(root, 'templates')):
				return os.path.join(root, 'templates')
		
		# Fallback to current directory
		return dev_template_dir

class MeinAntragApp(BaseTemplateResource):
	def __init__(self):
		# Setup Jinja2 template environment
		template_dir = self._get_template_dir()
		print(f"Using template directory: {template_dir}")
		self.jinja_env = Environment(loader=FileSystemLoader(template_dir))
	
	def on_get(self, req, resp):
		"""Serve the main page"""
		template = self.jinja_env.get_template('index.html')
		resp.content_type = 'text/html; charset=utf-8'
		resp.text = template.render(
			meta_title='MeinAntrag – Anträge an die Karlsruher Stadtverwaltung',
			meta_description='Erstelle einfach Vorlagen für Anfragen oder Anträge an die Karlsruher Stadtverwaltung zu deinem persönlichen Thema und schicke diese direkt an eine Stadtratsfraktion!',
			canonical_url=f"{SITE_BASE_URL}/"
		)

class ImpressumResource(BaseTemplateResource):
	def __init__(self):
		template_dir = self._get_template_dir()
		self.jinja_env = Environment(loader=FileSystemLoader(template_dir))
	
	def on_get(self, req, resp):
		"""Serve the Impressum page"""
		template = self.jinja_env.get_template('impressum.html')
		resp.content_type = 'text/html; charset=utf-8'
		resp.text = template.render(
			meta_title='Impressum – MeinAntrag',
			meta_description='Impressum für MeinAntrag.',
			canonical_url=f"{SITE_BASE_URL}/impressum",
			noindex=True
		)

class DatenschutzResource(BaseTemplateResource):
	def __init__(self):
		template_dir = self._get_template_dir()
		self.jinja_env = Environment(loader=FileSystemLoader(template_dir))
	
	def on_get(self, req, resp):
		"""Serve the Datenschutz page"""
		template = self.jinja_env.get_template('datenschutz.html')
		resp.content_type = 'text/html; charset=utf-8'
		resp.text = template.render(
			meta_title='Datenschutz – MeinAntrag',
			meta_description='Datenschutzerklärung für MeinAntrag. Keine Cookies, es werden nur Anfragen an die FragDenStaat-API gestellt.',
			canonical_url=f"{SITE_BASE_URL}/datenschutz",
			noindex=True
		)

class GenerateAntragResource:
	def __init__(self):
		# Initialize Gemini API
		api_key = os.environ.get('GOOGLE_GEMINI_API_KEY')
		if api_key:
			genai.configure(api_key=api_key)
			self.model = genai.GenerativeModel('gemini-flash-latest')
		else:
			self.model = None
	
	def _remove_markdown(self, text):
		"""Remove markdown formatting from text"""
		if not text:
			return text
		
		# Remove bold/italic markdown: **text** or *text* or __text__ or _text_
		text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
		text = re.sub(r'\*(.+?)\*', r'\1', text)
		text = re.sub(r'__(.+?)__', r'\1', text)
		text = re.sub(r'_(.+?)_', r'\1', text)
		
		# Remove heading markdown: /Heading or # Heading
		text = re.sub(r'^/\s*', '', text, flags=re.MULTILINE)
		text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
		
		# Remove other markdown elements
		text = re.sub(r'`(.+?)`', r'\1', text)  # Code
		text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # Links
		
		return text.strip()
	
	def _parse_gemini_response(self, text):
		"""Parse the response from Gemini into title, demand, and justification"""
		# Remove markdown formatting first
		text = self._remove_markdown(text)
		
		# Split by "Begründung/Sachverhalt" or similar patterns
		parts = re.split(r'(begründung|sachverhalt|begründung/sachverhalt)', text, maxsplit=1, flags=re.IGNORECASE)
		
		if len(parts) >= 3:
			# We have a split at "Begründung/Sachverhalt"
			before_justification = parts[0].strip()
			justification = parts[2].strip() if len(parts) > 2 else ""
			
			# Remove "Begründung/Sachverhalt" or "Sachverhalt" from the beginning of justification
			justification = re.sub(r'^(begründung\s*/?\s*sachverhalt|sachverhalt|begründung)\s*:?\s*\n?', '', justification, flags=re.IGNORECASE).strip()
		else:
			# Try to split by paragraphs
			paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
			if len(paragraphs) >= 3:
				before_justification = '\n\n'.join(paragraphs[:-1])
				justification = paragraphs[-1]
				# Remove "Begründung/Sachverhalt" or "Sachverhalt" from the beginning
				justification = re.sub(r'^(begründung\s*/?\s*sachverhalt|sachverhalt|begründung)\s*:?\s*\n?', '', justification, flags=re.IGNORECASE).strip()
			else:
				before_justification = text
				justification = ""
		
		# Extract title (first line or first paragraph)
		title_match = re.match(r'^(.+?)(?:\n\n|\n|$)', before_justification)
		if title_match:
			title = title_match.group(1).strip()
			demand = before_justification[len(title):].strip()
		else:
			lines = before_justification.split('\n', 1)
			title = lines[0].strip()
			demand = lines[1].strip() if len(lines) > 1 else ""
		
		# Remove title from demand if it's duplicated
		if demand.startswith(title):
			demand = demand[len(title):].strip()
		
		# Remove markdown from each part
		title = self._remove_markdown(title)
		demand = self._remove_markdown(demand)
		justification = self._remove_markdown(justification)
		
		# Final cleanup: remove any remaining "Sachverhalt" or "Begründung/Sachverhalt" at the start
		justification = re.sub(r'^(begründung\s*/?\s*sachverhalt|sachverhalt|begründung)\s*:?\s*\n?\s*', '', justification, flags=re.IGNORECASE).strip()
		
		return {
			'title': title,
			'demand': demand,
			'justification': justification
		}
	
	def on_post(self, req, resp):
		"""Generate text from user input using Gemini API"""
		try:
			if not self.model:
				resp.status = falcon.HTTP_500
				resp.content_type = 'application/json'
				resp.text = json.dumps({
					'success': False,
					'error': 'Gemini API key not configured'
				})
				return
			
			# Get form data - try multiple methods for Falcon compatibility
			anliegen = ''
			party_id = ''
			
			# Method 1: Try get_param (works for URL-encoded form data)
			anliegen = req.get_param('anliegen', default='') or ''
			party_id = req.get_param('party_id', default='') or ''
			
			# Method 2: If empty, try to read from stream and parse manually
			if not anliegen:
				try:
					# Read the raw body - use bounded_stream if available, otherwise stream
					stream = getattr(req, 'bounded_stream', req.stream)
					raw_body = stream.read().decode('utf-8')
					# Parse URL-encoded data manually
					parsed = parse_qs(raw_body)
					anliegen = parsed.get('anliegen', [''])[0]
					party_id = parsed.get('party_id', [''])[0]
				except Exception as e:
					# Log the exception for debugging
					print(f"Error parsing form data: {e}")
					pass
			
			# Remove any whitespace and check if actually empty
			anliegen = anliegen.strip() if anliegen else ''
			party_id = party_id.strip() if party_id else ''
			
			if not anliegen:
				resp.status = falcon.HTTP_400
				resp.content_type = 'application/json'
				resp.text = json.dumps({
					'success': False,
					'error': 'Anliegen-Feld ist erforderlich'
				})
				return
			
			# Create prompt for Gemini
			prompt = """Erzeuge aus dem folgenden Anliegen-Text je nach Anliegen eine Anfrage oder einen Antrag an die Karlsruher Stadtverwaltung im Namen einer Stadtratsfraktion. 

Der Antrag soll im sachlichen, offiziellen Ton einer Fraktion verfasst sein - KEINE persönliche Anrede, KEINE "ich" oder "wir" Formulierungen. Verwende die dritte Person oder Passiv-Formulierungen.

Struktur:
- Die erste Zeile ist der Antragstitel. Der Titel soll PRÄGNANT, EINFACH und EINPRÄGSAM sein - maximal 8-10 Wörter. Vermeide komplizierte Formulierungen, technische Fachbegriffe oder zu lange Titel. Der Titel soll eine gute Außenwirkung haben und das Anliegen klar und verständlich kommunizieren. Beispiele für gute Titel: "Nachtabsenkung der öffentlichen Straßenbeleuchtung", "Vielfalt in Bewegung – Kulturelle Begleitmaßnahmen World Games 2029", "Prüfung digitaler Zahlungsdienstleister und WERO-Alternative"
- Der zweite Absatz ist der Forderungsteil. Hier können nach einem kurzen Satz auch Stichpunkte verwendet werden, wenn dies sinnvoll ist.
- Der letzte Teil ist Begründung/Sachverhalt (ohne diesen Titel im Text)

WICHTIG: 
- Verwende KEINE Markdown-Formatierung. Keine **fett**, keine *kursiv*, keine /Überschriften, keine # Hashtags, keine Links oder andere Formatierung. 
- Schreibe nur reinen Text ohne jegliche Markdown-Syntax.
- Sachlicher, offizieller Ton einer Fraktion, keine persönlichen Formulierungen.
- Der Antragstitel muss prägnant, einfach verständlich und einprägsam sein - keine komplizierten Formulierungen!

"""
			prompt += anliegen
			
			# Call Gemini API
			response = self.model.generate_content(prompt)
			generated_text = response.text
			
			# Parse the response
			parsed = self._parse_gemini_response(generated_text)
			
			# Generate email text
			email_prompt = f"""Erstelle einen kurzen, höflichen E-Mail-Text in der ERSTEN PERSON (ich/wir) für eine Fraktion an eine andere Fraktion. 
Die E-Mail soll:
- Mit "Guten Tag," beginnen
- Das Anliegen kurz in der ersten Person erklären (basierend auf: {anliegen})
- Erwähnen, dass eine Antragsvorlage im Anhang beigefügt ist
- Mit "Mit freundlichen Grüßen," enden

Der Text soll sachlich, höflich und kurz sein (2-3 Sätze zwischen Begrüßung und Grußformel). Verwende KEINE Markdown-Formatierung. Schreibe in der ERSTEN PERSON (z.B. "ich möchte", "wir bitten", "ich habe").

Anliegen: {anliegen}
"""
			
			email_response = self.model.generate_content(email_prompt)
			email_text = self._remove_markdown(email_response.text)
			
			# Ensure proper format - clean up and ensure structure
			email_text = email_text.strip()
			
			# Ensure it starts with "Guten Tag,"
			if not email_text.startswith('Guten Tag'):
				# Remove any existing greeting
				email_text = re.sub(r'^(Guten Tag[,\s]*|Hallo[,\s]*|Sehr geehrte[^,]*,\s*)', '', email_text, flags=re.IGNORECASE)
				email_text = 'Guten Tag,\n\n' + email_text.strip()
			
			# Ensure it ends with "Mit freundlichen Grüßen,"
			if 'Mit freundlichen Grüßen' not in email_text:
				email_text += '\n\nMit freundlichen Grüßen,'
			else:
				# Make sure it's properly formatted
				if not email_text.rstrip().endswith('Mit freundlichen Grüßen,'):
					# Remove any existing closing and add proper one
					email_text = re.sub(r'\s*Mit freundlichen Grüßen[,\s]*$', '', email_text, flags=re.IGNORECASE)
					email_text = email_text.rstrip() + '\n\nMit freundlichen Grüßen,'
			
			# Return JSON with the generated text parts
			resp.content_type = 'application/json'
			resp.text = json.dumps({
				'success': True,
				'title': parsed['title'],
				'demand': parsed['demand'],
				'justification': parsed['justification'],
				'email_body': email_text,
				'party_name': party_id if party_id else ""
			})
			
		except Exception as e:
			import traceback
			traceback.print_exc()
			resp.status = falcon.HTTP_500
			resp.content_type = 'application/json'
			resp.text = json.dumps({
				'success': False,
				'error': str(e)
			})

class GenerateWordResource:
	def __init__(self):
		# Get template path
		script_dir = os.path.dirname(os.path.abspath(__file__))
		self.template_path = os.path.join(script_dir, 'assets', 'antrag_vorlage.docx')
		# Fallback if not in assets
		if not os.path.exists(self.template_path):
			assets_dir = os.path.join(script_dir, '..', 'assets')
			self.template_path = os.path.join(assets_dir, 'antrag_vorlage.docx')
	
	def _generate_word(self, title, demand, justification, party_name=""):
		"""Generate a Word document using the template"""
		# Load template
		if os.path.exists(self.template_path):
			doc = Document(self.template_path)
		else:
			# Fallback: create new document if template not found
			doc = Document()
		
		# Get current date in DD.MM.YYYY format
		current_date = datetime.now().strftime("%d.%m.%Y")
		
		# Use demand directly without heading
		antragtext = demand
		
		# Replace placeholders in all paragraphs
		for paragraph in doc.paragraphs:
			full_text = paragraph.text
			if not full_text:
				continue
			
			# Replace FRAKTION
			if 'FRAKTION' in full_text:
				for run in paragraph.runs:
					if 'FRAKTION' in run.text:
						run.text = run.text.replace('FRAKTION', party_name)
			
			# Replace XX.XX.XXXX with current date
			if 'XX.XX.XXXX' in full_text:
				for run in paragraph.runs:
					if 'XX.XX.XXXX' in run.text:
						run.text = run.text.replace('XX.XX.XXXX', current_date)
			
			# Replace ANTRAGSTITEL (bold)
			if 'ANTRAGSTITEL' in full_text:
				paragraph.clear()
				run = paragraph.add_run(title)
				run.bold = True
			
			# Replace ANTRAGSTEXT
			if 'ANTRAGSTEXT' in full_text:
				paragraph.clear()
				lines = antragtext.split('\n')
				for i, line in enumerate(lines):
					if line.strip():
						paragraph.add_run(line.strip())
						if i < len(lines) - 1:
							paragraph.add_run('\n')
			
			# Replace BEGRÜNDUNGSTEXT
			if 'BEGRÜNDUNGSTEXT' in full_text:
				paragraph.clear()
				lines = justification.split('\n')
				for i, line in enumerate(lines):
					if line.strip():
						paragraph.add_run(line.strip())
						if i < len(lines) - 1:
							paragraph.add_run('\n')
		
		# Check text boxes (shapes) for placeholders
		# Text boxes are stored in the document's part relationships
		try:
			# Access document part to search for text boxes
			document_part = doc.part
			from docx.oxml.ns import qn
			
			# Search for FRAKTION in text boxes
			# Text boxes are in w:txbxContent elements within w:p (paragraphs)
			# We need to search the entire XML tree
			def replace_in_element(element, search_text, replace_text):
				"""Recursively replace text in XML elements"""
				if element.text and search_text in element.text:
					element.text = element.text.replace(search_text, replace_text)
				if element.tail and search_text in element.tail:
					element.tail = element.tail.replace(search_text, replace_text)
				for child in element:
					replace_in_element(child, search_text, replace_text)
			
			# Search in main document body
			if party_name:
				replace_in_element(document_part.element, 'FRAKTION', party_name)
			
			# Also search in header and footer parts
			for rel in document_part.rels.values():
				if 'header' in rel.target_ref or 'footer' in rel.target_ref:
					try:
						header_footer_part = rel.target_part
						if party_name:
							replace_in_element(header_footer_part.element, 'FRAKTION', party_name)
					except Exception:
						pass
		except Exception as e:
			# If text box access fails, continue with other replacements
			print(f"Warning: Could not replace in text boxes: {e}")
			pass
		
		# Also check tables for placeholders
		for table in doc.tables:
			for row in table.rows:
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						full_text = paragraph.text
						if not full_text:
							continue
						
						if party_name and 'FRAKTION' in full_text:
							for run in paragraph.runs:
								if 'FRAKTION' in run.text:
									run.text = run.text.replace('FRAKTION', party_name)
						
						if 'XX.XX.XXXX' in full_text:
							for run in paragraph.runs:
								if 'XX.XX.XXXX' in run.text:
									run.text = run.text.replace('XX.XX.XXXX', current_date)
						
						if 'ANTRAGSTITEL' in full_text:
							paragraph.clear()
							run = paragraph.add_run(title)
							run.bold = True
						
						if 'ANTRAGSTEXT' in full_text:
							paragraph.clear()
							lines = antragtext.split('\n')
							for i, line in enumerate(lines):
								if line.strip():
									paragraph.add_run(line.strip())
									if i < len(lines) - 1:
										paragraph.add_run('\n')
						
						if 'BEGRÜNDUNGSTEXT' in full_text:
							paragraph.clear()
							lines = justification.split('\n')
							for i, line in enumerate(lines):
								if line.strip():
									paragraph.add_run(line.strip())
									if i < len(lines) - 1:
										paragraph.add_run('\n')
		
		# Save to buffer
		buffer = BytesIO()
		doc.save(buffer)
		buffer.seek(0)
		return buffer
	
	def on_post(self, req, resp):
		"""Generate Word document from form data"""
		try:
			if not DOCX_AVAILABLE:
				resp.status = falcon.HTTP_500
				resp.content_type = 'application/json'
				resp.text = json.dumps({
					'success': False,
					'error': 'python-docx not installed'
				})
				return
			
			# Get form data
			title = req.get_param('title', default='') or ''
			demand = req.get_param('demand', default='') or ''
			justification = req.get_param('justification', default='') or ''
			party_name = req.get_param('party_name', default='') or ''
			
			# If empty, try to read from stream
			if not title:
				try:
					stream = getattr(req, 'bounded_stream', req.stream)
					raw_body = stream.read().decode('utf-8')
					parsed = parse_qs(raw_body)
					title = parsed.get('title', [''])[0]
					demand = parsed.get('demand', [''])[0]
					justification = parsed.get('justification', [''])[0]
					party_name = parsed.get('party_name', [''])[0]
				except Exception:
					pass
			
			# Generate Word document
			word_buffer = self._generate_word(title, demand, justification, party_name)
			
			# Return Word document
			resp.content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
			resp.set_header('Content-Disposition', 'attachment; filename="antrag.docx"')
			resp.data = word_buffer.read()
			
		except Exception as e:
			import traceback
			traceback.print_exc()
			resp.status = falcon.HTTP_500
			resp.content_type = 'application/json'
			resp.text = json.dumps({
				'success': False,
				'error': str(e)
			})

class RobotsResource:
	def on_get(self, req, resp):
		resp.content_type = 'text/plain; charset=utf-8'
		resp.text = f"""User-agent: *
Allow: /
Sitemap: {SITE_BASE_URL}/sitemap.xml
"""

class SitemapResource:
	def on_get(self, req, resp):
		resp.content_type = 'application/xml; charset=utf-8'
		resp.text = f"""<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
  <url><loc>{SITE_BASE_URL}/</loc></url>
  <url><loc>{SITE_BASE_URL}/impressum</loc></url>
  <url><loc>{SITE_BASE_URL}/datenschutz</loc></url>
</urlset>
"""

# Create Falcon application
app = falcon.App()

# Discover static assets directory
STATIC_DIR = os.environ.get('MEINANTRAG_STATIC_DIR')
if not STATIC_DIR:
	# Prefer local assets folder in development (relative to this file)
	script_dir = os.path.dirname(os.path.abspath(__file__))
	candidate = os.path.join(script_dir, 'assets')
	if os.path.isdir(candidate):
		STATIC_DIR = candidate
	else:
		# Try current working directory (useful when running packaged binary from project root)
		cwd_candidate = os.path.join(os.getcwd(), 'assets')
		if os.path.isdir(cwd_candidate):
			STATIC_DIR = cwd_candidate
		else:
			# Fallback to packaged location under share
			STATIC_DIR = os.path.join(script_dir, '..', 'share', 'meinantrag', 'assets')

# Add routes
meinantrag = MeinAntragApp()
impressum = ImpressumResource()
datenschutz = DatenschutzResource()
generate_antrag = GenerateAntragResource()
generate_word = GenerateWordResource()
robots = RobotsResource()
sitemap = SitemapResource()

app.add_route('/', meinantrag)
app.add_route('/impressum', impressum)
app.add_route('/datenschutz', datenschutz)
app.add_route('/api/generate-antrag', generate_antrag)
app.add_route('/api/generate-word', generate_word)
app.add_route('/robots.txt', robots)
app.add_route('/sitemap.xml', sitemap)

# Static file route
if STATIC_DIR and os.path.isdir(STATIC_DIR):
	app.add_static_route('/static', STATIC_DIR)

if __name__ == '__main__':
	import wsgiref.simple_server
	
	print("Starting MeinAntrag web application...")
	print("Open your browser and navigate to: http://localhost:8000")
	print(f"Serving static assets from: {STATIC_DIR}")
	
	httpd = wsgiref.simple_server.make_server('localhost', 8000, app)
	httpd.serve_forever()
